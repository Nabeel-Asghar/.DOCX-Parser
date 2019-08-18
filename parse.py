import docx
from copy import deepcopy
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.document import Document as _Document
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import extract

def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def iter_hyperlink_rels(rels):
    dict = {'key' : 'value'}
    for rel in rels:
        if rels[rel].reltype == RT.HYPERLINK:
            if "/projects.webappsec.org" in rels[rel]._target:
                x = rels[rel]._rId[-2:]
                dict[x] = rels[rel]._target


    wasc = []
    for key, value in sorted(dict.items()): # Note the () after items!
        wasc.append(value[30:])
    return(wasc)

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

#Get number of errors from detailed report
def count(detailed):

    doc2 = docx.Document(detailed)

    tcount = 0
    n=1
    titles = []
    severity = []

    while 1<2:
        try:
            doc2.tables[1].cell(n,1).text
            titles.append(doc2.tables[1].cell(n,1).text)
            severity.append("High")
            tcount+=1
            n+=1
        except:
            break

    n=1
    while 1<2:
        try:
            doc2.tables[2].cell(n,1).text
            titles.append(doc2.tables[2].cell(n,1).text)
            severity.append("Medium")
            tcount+=1
            n+=1
        except:
            break

    n=1
    while 1<2:
        try:
            doc2.tables[3].cell(n,1).text
            titles.append(doc2.tables[3].cell(n,1).text)
            severity.append("Low")
            tcount+=1
            n+=1
        except:
            break

    n=1
    while 1<2:
        try:
            doc2.tables[4].cell(n,1).text
            titles.append(doc2.tables[4].cell(n,1).text)
            severity.append("Low")
            tcount+=1
            n+=1
        except:
            break

    return(tcount,titles,severity)

#Copy images
def images(ourDict):

    table = []

    for key,value in ourDict.items():
        for i in range(value):
            table.append(key)

    document = docx.Document('Final-Report.docx')

    for i in range(len(table)):
        x = table[i] + 2
        cell = document.tables[x].cell(11,6)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        try:
            run.add_picture('extraction\word\media\image%s.png' % (7+i),  width = 4500000)
        except:
            break
    document.save("Final-Report.docx")

#Function to copy table
def copy_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    new_tbl = deepcopy(tbl)
    p.addnext(new_tbl)

#Get hyperlinks - currently not working
def links(detailed):
    document = docx.Document(detailed)
    rels = document.part.rels
    wasc = iter_hyperlink_rels(rels)
    return(wasc)

#Fill in first page of report from testplan
def testplan(testpln):
    doc = docx.Document(testpln)

    projectid = doc.tables[1].cell(1,4).text
    appid = doc.tables[1].cell(3,4).text
    # tester = doc.tables[1].cell(?,?).text
    dateissue = doc.tables[1].cell(15,0).text
    scope = doc.tables[1].cell(9,4).text
    ateam = doc.tables[1].cell(14,4).text
    architect = doc.tables[1].cell(1,0).text
    dmanager = doc.tables[1].cell(5,0).text
    techlead = doc.tables[1].cell(11,0).text
    # testtype = doc.tables[1].cell(?,?).text
    urls = doc.tables[1].cell(19,1).text
    # ip = doc.tables[0].cell(14,0).text
    comments = doc.tables[1].cell(29,0).text


    #Write to final report
    doc3 = docx.Document('Final-Report.docx')

    doc3.tables[0].cell(2,0).text = projectid
    doc3.tables[0].cell(4,0).text = appid
    #doc3.tables[0].cell(6,0).text = tester
    doc3.tables[0].cell(8,0).text = dateissue

    doc3.tables[0].cell(2,3).text = ateam
    doc3.tables[0].cell(4,3).text = architect
    doc3.tables[0].cell(6,3).text = dmanager
    doc3.tables[0].cell(8,3).text = techlead

    doc3.tables[0].cell(12,0).text = urls
    #doc3.tables[0].cell(14,0).text = ip
    doc3.tables[0].cell(16,0).text = comments
    doc3.save('Final-Report.docx')

#Get input in a formatted way
def getFormat(detailed):

    limit,list,severity = count(detailed)
    limit += 7
    tablecount = 0

    document = docx.Document(detailed)

    sections = document.sections

    list = []

    for block in iter_block_items(document):
         x = block if isinstance(block, Paragraph) else '<table>'
         list.append(x)

    formatList = []

    for object in list:
        if tablecount > limit:
            break

        elif(object == '<table>'):
            tablecount+=1
            if  tablecount > 7 and tablecount<=limit:
                formatList.append("\n\nStartOfTable\n\n")
                #print("\n\n----------------------------Table----------------------------\n\n")
        elif tablecount >= 7 and tablecount<=limit:
            formatList.append(object.text)
            #print(object.text)

    return(formatList)

#Puts as many tables as errors
def getTables(count):
    document = docx.Document('Final-Report.docx')
    for para in document.paragraphs:
        if para.text == "Do not delete":
            variable = para

    table = document.tables[2]
    paragraph = variable

    for i in range(count-1):
        copy_table_after(table, paragraph)
        document.add_page_break()
        paragraph = document.add_paragraph('Do not delete')
        document.save('Final-Report.docx')

#Input title, description,severity,occurences into final report
def titleAndDesc(detailed):


    limit,titles,severity = count(detailed)
    wasc = links(detailed)

    document = docx.Document('Final-Report.docx')

    big = []
    urls = []
    reason = []
    bigresponse = []
    whichtable = -1
    ourList = getFormat(detailed)
    i=0
    tablenum = 2
    occurences = []
    image = []
    dict = {}

    for i in range(len(ourList)):

        if ourList[i] == "Advisory & Fix Recommendation":
            whichtable+=1
            i+=1
            small = []
            while ourList[i] != "Affected URL":
                small.append(ourList[i])
                if ourList[i] != "" :
                    small.append("\n")
                i+=1
            big.append(small)

        if ourList[i] == "Affected URL":
            i+=1
            url = []
            urlnum = 0

            while "://" in ourList[i] and "Issue" not in ourList[i]:
                url.append(ourList[i])
                url.append("\n")
                urlnum+=1
                i+=1
            occurences.append(urlnum)
            urls.append(url)

        if "Variant Reasoning" in ourList[i]:

            reason.append(whichtable)
            reason.append(ourList[i])

        if "Variant Request Response" in ourList[i]:
            response = []
            try:
                while ourList[i] != "Advisory & Fix Recommendation" and "Issue" not in ourList[i]:
                    response.append(ourList[i])
                    response.append("\n")
                    i+=1
            except IndexError:
                pass
            bigresponse.append(whichtable)
            bigresponse.append(response)
        try:
            if "Issue" in ourList[i] and "/" in ourList[i]:
                temp = 1
                i+=1
                try:
                    while "StartOfTable" not in ourList[i]:
                        if "Issue" in ourList[i] and "/" in ourList[i]:
                            temp+=1
                            i+=1
                        else:
                            i+=1
                except IndexError:
                    pass

                if whichtable not in dict:
                    dict[whichtable] = temp
        except IndexError:
            pass

    reason = list(dict.fromkeys(reason))

    style = document.styles['Heading 2']
    font = style.font
    font.name = 'Arial'

    #URLS
    realurls = []
    realoccurences = []
    for x in range(len(urls)):
        if x % 2 == 0:
            realurls.append(urls[x])
            realoccurences.append(occurences[x])

    for i in range(len(realurls)):
        document.tables[i+2].cell(12,6).text = "%s" % realoccurences[i]
        document.tables[i+2].cell(13,0).text = realurls[i]

    for i in range(limit):
        document.tables[i+2].cell(7,0).text = big[i]
        document.tables[i+2].cell(6,6).text = severity[i]
        document.tables[i+2].cell(4,6).text = "Open"
        document.tables[i+2].cell(0,0).add_paragraph("Observation %s: %s" % (i+1,titles[i]), style='Heading 2')
        document.tables[i+2].cell(2,6).text = wasc[i]

    for i in range(len(reason)):
        if i % 2 == 0:
            x = reason[i] + 2
            try:
                document.tables[x].cell(9,6).text = reason[i+1]
            except:
                pass

    for i in range(len(bigresponse)):
        if i % 2 == 0:
            x = bigresponse[i] + 2
            #document.tables[x].cell(10,6).text = bigresponse[i+1]
            document.tables[x].cell(10,6).add_paragraph("%s" % (bigresponse[i+1]),  style='Style Variant')

    document.save('Final-Report.docx')
    return(dict)

#main
def main():
    testpln = str(input("Enter test plan name including extension (.docx): "))
    detailed = str(input("Enter detailed report name including extension (.docx): "))
    print("\n")
    extract.extraction(detailed)
    testplan(testpln)
    print("Completed Test Plan portion")
    limit,titles,severity = count(detailed)
    getTables(limit)
    ourDict = titleAndDesc(detailed)

    images(ourDict)
    print("Completed Detailed Report portion")
    document = docx.Document('Final-Report.docx')

    for para in document.paragraphs:
        if para.text == "Do not delete":

            delete_paragraph(para)


    document.save('Final-Report.docx')
    print("Condensed report has been completed.")


if __name__== "__main__":
  main()
